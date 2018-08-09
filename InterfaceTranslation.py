# -*- coding: utf-8 -*-

import sys
import os
import docx
# 后台接口文档翻译
# 将数据库文档中 字段 解释 加入到接口文档中的字段解释

# 记得在 接口文档末尾加上 终止符 @wangbinbinend
# 要翻译的接口文档
readPath = r'/Users/sky-fish/Desktop/后台接口文\ 文档/接口文档 '
# 完成翻译之后的文档
outPath = r'接口文档-已完成'
# 数据库文档
dateFilePath = r'/Users/sky-fish/Desktop/后台接口文\ 文档/数据库设计.docx'
# 文档终止符，要手动把 @wangbinbinend 加入到接口文档结尾处。
END_FLAG = '@wangbinbinend'

# 获取接口中所有 句子， 所有字段集合
def getStrLists(f,fieldList,lineList):
    while True:
        # 逐行读取文件
        line = f.readline()
        (isSucc, feild) = interceptionStrBeforeColon(line)
        if isSucc == True:
            fieldList.append(feild)
        else:
            # 为了保持 fieldList 和 lineList 的元素个数想等，当 isSucc 为 False 时，给fieldList 中添加一个占位元素
            fieldList.append('占位元素')
        # lineList 装入 原来的接口文档中的句子
        lineList.append(line)
        # 当到了终止符 @wangbinbinend 时，说明接口文档中所有的句子都处理结束，退出循环
        if(END_FLAG in line):
            break

    fieldListCount = len(fieldList)
    lineListCount = len(lineList)
    # fieldListCount 和 lineListCount的元素个数肯定是想等的
    if fieldListCount != lineListCount:
        raise RuntimeError('fieldList 和 lineList元素个数不同，这里要求这两个 list 必须有相同的 list')
    return (fieldList,lineList)

# 获取接口中所有字段
def readField(fileName):
    # 打开文件
    f = open(fileName)
    # 字段集合,用来和数据库字段做比较
    fieldList = []
    # 原文件每行 文字集合
    lineList = []
    return getStrLists(f,fieldList,lineList)


# 数据库中字段 去掉下划线如 ：my_User_Direction 变为 myUserDirection
def removeCharacter(string,character):
    stringList = list(string)
    # 记录有多少个 字符 '_' ,初始化为零
    cout = 0
    for element in stringList:
        if element == '_':
            cout = cout + 1
    # 把字符串中 所有的 下划线都删除掉
    while True:
        if cout <= 0:
            break
        stringList.remove(character)
        cout = cout - 1

    return "".join(stringList)


# 读取word 数据库文档, 得到 数据库 字段list
def readWordFile(wordPath):
    # 获取文档对象
    file = docx.Document(wordPath)
    # print("段落数:" + str(len(file.paragraphs)))  # 段落数为13，每个回车隔离一段
    # 输出每一段的内容
    # for para in file.paragraphs:
    #     print(para.text)
    # 输出段落编号及段落内容
    # for i in range(len(file.paragraphs)):
    #     print("第" + str(i) + "段的内容是：" + file.paragraphs[i].text)
    # 用来存 我们要的字段和字段解释
    feild_explain = []
    for table in file.tables:
        for row in table.rows:
            feild_cell_row = []
            for i,cell in enumerate(row.cells):
                # 我们只要 第0列 和 第一列的内容
                if i == 0 or i == 1:
                    cellText = cell.text
                    # 下面都是在 筛选 符合条件的内容
                    cellText = str(cellText).strip()
                    if len(cellText) <= 0:
                        continue
                    if cellText is None:
                        continue
                    if '字段说明' in cellText:
                        continue
                    if '字段名' in cellText:
                        continue
                    if cellText in ['']:
                        continue
                    # 剔除 字段中的下划线
                    cellText = removeCharacter(cellText, '_')
                    # 转大写
                    cellText = cellText.upper()
                    feild_cell_row.append(cellText)
                    # print('cell.text = %s' % (cell.text))
            if len(feild_cell_row) <= 1:
                continue
            feild_explain.append(feild_cell_row)
    # print('feild_explain = %s' % (feild_explain))
    return  feild_explain

# 截取字符串中我们想要的字段 如把接口文档中的字符串: '"parentLabel": "销售管理"' 截取得到 "parentLabel"
def interceptionStrBeforeColon(line):
    # 判断这行文字中是否有 冒号:, 接口文档一般都是: "username":"1" 这种形式，其中必有 冒号
    if ':' in line:
        # 判断字符串中有无 引号", 接口正式内容肯定有引号
        if "\"" in line:
            if ':' not in line:
                return (True,line)
            # 使用"来切割字符 如 "username":"1" 切割后就是
            field = line.split(r':')  # 分割字符串 得到字段
            # 冒号前面的字段
            field = field[0]
            # 删除所有空白字符 \n  \t  \r  '' 等
            field = field.strip()
            # 去除 引号 如 ： "user" 变为 user
            field = field.split('\"')
            field = field[1]
            # 转大写
            field = field.upper()
            #print('field = %s',field)
            return (True,field)# 得到想要的字段
        else:
            return (False,line) # 未得到我们要的字段
    else:
        return (False,line) # 未得到我们要的字段


# 对比接口文档字段 数据库文档字段，把字段解释加入接口list中
def compareTwoList(lineList,feildList,feild_explain):
    for index, feild in enumerate(feildList):
        for feild_value in feild_explain:
            feild_value_01 = feild_value[1]
            # 如果 feild_value_01 == feild 说明找到了我们想要字段的解释
            if feild_value_01 == feild:
                # 用来代替接口文档的新行
                temp = []
                # 取出接口文档中的 第 index 行
                temp_part_one = lineList[index]
                # 把这一行的换行符去掉，方便在后面添加字符
                temp_part_one = temp_part_one.replace('\n', '')
                # 把修改过的旧行装入到 temp
                temp.append(temp_part_one)
                feild_value00 = str(feild_value[0])
                # 加入换行符
                feild_value00 = feild_value00 + '\n'
                # 把字段解释装入到 temp
                temp.append(feild_value00)
                # 将list类型转为 str 类型，list 元素用 // 连接
                temp = r'//'.join(temp)
                print("temp = %s" % (temp))
                # 用新接口行，替换旧行
                lineList[index] = temp
                print('修改后的 lineList 第 %d 个元素: %s' % (index, lineList[index]))
                break
    # 打印测试
    print('lineList = %s' % (lineList))
    return lineList;


# 输入多行文字，写入指定文件并保存到指定文件夹
def writeFile(filename,content):
    fopen = open(filename, 'w')
    for line in content:
        fopen.write(line)
    fopen.close()


# 测试方法： 将list 转换为 str
def joinList(feild_explain):
    feilds = []
    for element in feild_explain:
        count = len(element)
        if count<= 0:
            continue
        elementStr = str(element[count - 1])
        if len(elementStr) <= 0:
            continue
        elementStrList = list(elementStr)
        elementStrListCount = len(elementStrList)
        elementStrList.insert(elementStrListCount,'\n')
        elementStr = ''.join(elementStrList)
        element[count - 1] = elementStr
        element = '------'.join(element)
        feilds.append(element)
    return feilds


# 测试方法:将 list 写入文件，用来测试
def writeList(listName,fileName):
    fileName = open(fileName, 'w')
    for element in listName:
        element_list = list(element)
        element_list.append('\n')
        element = ''.join(element_list)
        fileName.write(element)
    fileName.close()


def main():
    # 读取数据库文档中字段及其翻译 得到一个二维 list feild_explain
    feild_explain = readWordFile(dateFilePath)
    # 读取接口文档中的句子，fieldList 字段list， lineList 句子 list
    (fieldList, lineList) = readField(readPath)
    # 通过对比，把接口文档中的句子 lineList 更改为已经翻译之后的句子
    lineList = compareTwoList(lineList,fieldList, feild_explain)
    # 将翻译后的句子，写入文档 outpath
    writeFile(outPath, lineList)


if __name__ == "__main__":
    main()
